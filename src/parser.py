"""
Resume parser module.
Handles text extraction and cleaning from various resume formats.
"""
import re
import pandas as pd


def clean_text(text: str) -> str:
    """Clean and normalize resume text."""
    if not isinstance(text, str):
        return ""
    
    # Convert to lowercase
    text = text.lower()
    
    # Remove extra whitespace
    text = re.sub(r'\s+', ' ', text)
    
    # Remove special characters but keep basic punctuation
    text = re.sub(r'[^\w\s.,;:()\-/]', '', text)
    
    return text.strip()


def load_resumes_from_csv(filepath: str) -> pd.DataFrame:
    """
    Load resumes from CSV file.
    Expected columns: candidate_id, name, resume_text
    """
    df = pd.read_csv(filepath)
    
    # Validate required columns
    required_cols = ['candidate_id', 'name', 'resume_text']
    missing = set(required_cols) - set(df.columns)
    if missing:
        raise ValueError(f"Missing required columns: {missing}")
    
    # Clean resume text
    df['resume_text_clean'] = df['resume_text'].apply(clean_text)
    
    return df


def parse_pdf_resume(file_object) -> str:
    """
    Parse PDF resume to extract text.
    
    Args:
        file_object: File-like object or path string
    
    Returns:
        Extracted text from PDF
    """
    try:
        import PyPDF2
        
        # Handle both file paths and file objects
        if isinstance(file_object, str):
            with open(file_object, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text()
            return text
        else:
            # File object (e.g., from Streamlit file_uploader)
            reader = PyPDF2.PdfReader(file_object)
            text = ""
            for page in reader.pages:
                text += page.extract_text()
            return text
    except PyPDF2.errors.PdfReadError as e:
        if "EOF marker not found" in str(e):
            raise ValueError("This PDF file is corrupted or invalid. Please upload a valid PDF file.")
        else:
            raise ValueError(f"Cannot read PDF: {str(e)}")
    except ImportError:
        raise ImportError(
            "PDF parsing requires PyPDF2. Install with: pip install PyPDF2"
        )
    except Exception as e:
        raise ValueError(f"Error parsing PDF: {str(e)}")


def parse_docx_resume(file_object) -> str:
    """
    Parse DOCX resume to extract text.
    
    Args:
        file_object: File-like object or path string
    
    Returns:
        Extracted text from DOCX
    """
    try:
        import docx
        
        # Handle both file paths and file objects
        if isinstance(file_object, str):
            doc = docx.Document(file_object)
        else:
            # File object (e.g., from Streamlit file_uploader)
            doc = docx.Document(file_object)
        
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    except ImportError:
        raise ImportError(
            "DOCX parsing requires python-docx. Install with: pip install python-docx"
        )


def parse_resume_file(file_object, file_name: str) -> str:
    """
    Master parsing function that handles multiple file types.
    
    Args:
        file_object: File-like object or path string
        file_name: Name of the file (used to determine extension)
    
    Returns:
        Extracted text from the file
    
    Raises:
        ValueError: If file type is not supported
    """
    # Get file extension
    file_ext = file_name.lower().split('.')[-1] if '.' in file_name else ''
    
    if file_ext == 'pdf':
        return parse_pdf_resume(file_object)
    elif file_ext in ['docx', 'doc']:
        return parse_docx_resume(file_object)
    elif file_ext == 'txt':
        # Handle text files
        if isinstance(file_object, str):
            with open(file_object, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        else:
            # File object from Streamlit
            return file_object.read().decode('utf-8', errors='ignore')
    else:
        raise ValueError(f"Unsupported file type: .{file_ext}. Supported types: pdf, docx, txt")


if __name__ == "__main__":
    # Test the parser
    sample_text = """
    John Doe - Software Engineer
    Experience: 5 years in Python, Django, AWS, Machine Learning
    Skills: Data Science, Docker, Kubernetes, React
    """
    
    cleaned = clean_text(sample_text)
    print("Cleaned text:", cleaned)
